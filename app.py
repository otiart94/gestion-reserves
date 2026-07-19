import streamlit as st
import base64
import json
import io
import os
import time
import datetime
from xhtml2pdf import pisa
from docx import Document
from PIL import Image, ImageOps, ImageDraw

# Configuration de la page Streamlit (Doit être la toute première commande Streamlit)
st.set_page_config(page_title="Générateur de Levées de Réserves", page_icon="🏗️", layout="wide")

import folium
from streamlit_folium import st_folium
import fitz  # PyMuPDF

st.title("🏗️ Assistant de Gestion & Levées de Réserves")
st.write("Remplissez les informations, placez vos réserves sur le plan (PDF/Image), et compilez votre rapport au format HTML/PDF.")

# --- INITIALISATION DE L'ÉTAT ---
if "tasks" not in st.session_state:
    st.session_state.tasks = []

if "loaded_json_file_name" not in st.session_state:
    st.session_state.loaded_json_file_name = None

# Nouveau dictionnaire pour stocker les plans par étage (clé: "BATIMENT - ETAGE", valeur: data_url de l'image)
if "floor_plans" not in st.session_state:
    st.session_state.floor_plans = {}

# --- FONCTIONS GLOBALES ---
def priority_weight(prio_str):
    """Donne un poids au tri de la priorité (Haute en premier)"""
    p = prio_str.lower() if prio_str else ""
    if "haute" in p: return 0
    if "basse" in p: return 2
    return 1

# --- FONCTIONS DE GESTION DES PHOTOS AVEC COMPRESSION ET RÉDUCTION À 800PX ---
def update_photo_callback(task_id, img_idx, key_uploader):
    """Gère le téléversement et compresse drastiquement l'image à 800px max pour éviter de saturer la RAM"""
    file = st.session_state[key_uploader]

    task_idx = next((index for index, t in enumerate(st.session_state.tasks) if t["id"] == task_id), None)

    if task_idx is not None and file is not None:
        while len(st.session_state.tasks[task_idx]["images"]) <= img_idx:
            st.session_state.tasks[task_idx]["images"].append({"b64": "", "name": ""})

        try:
            raw_image = Image.open(file)
            fixed_image = ImageOps.exif_transpose(raw_image)

            if fixed_image.mode != "RGB":
                fixed_image = fixed_image.convert("RGB")

            # --- REDIMENSIONNEMENT À 800PX ---
            max_size = 800
            fixed_image.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)

            buffered = io.BytesIO()
            fixed_image.save(buffered, format="JPEG", quality=60, optimize=True)

            st.session_state.tasks[task_idx]["images"][img_idx] = {
                "b64": base64.b64encode(buffered.getvalue()).decode("utf-8"),
                "name": file.name
            }
            st.session_state.tasks[task_idx]["_img_version"] = st.session_state.tasks[task_idx].get("_img_version", 0) + 1
        except Exception as e:
            st.error(f"Erreur lors du traitement de l'image : {e}")

def delete_photo_callback(task_id, img_idx):
    """Supprime proprement une photo de l'état"""
    task_idx = next((index for index, t in enumerate(st.session_state.tasks) if t["id"] == task_id), None)
    if task_idx is not None:
        if img_idx < len(st.session_state.tasks[task_idx]["images"]):
            st.session_state.tasks[task_idx]["images"][img_idx] = {"b64": "", "name": ""}
            st.session_state.tasks[task_idx]["_img_version"] = st.session_state.tasks[task_idx].get("_img_version", 0) + 1

# --- CACHE DE SÉRIALISATION JSON PAR RÉSERVE ---
if "task_json_cache" not in st.session_state:
    st.session_state.task_json_cache = {}

def get_task_export_json(task):
    cache = st.session_state.task_json_cache
    t_id = task["id"]
    img_version = task.get("_img_version", 0)
    cache_key = (
        task.get("num_reserve", ""), task.get("batiment", ""), task.get("etage", ""), task.get("priorite", ""), 
        task.get("categorie", ""), task.get("titre", ""), task.get("details", ""), task.get("commentaire", ""), 
        task.get("status", ""), task.get("status_logistique", ""), task.get("entreprise", ""), task.get("date_livraison", ""), task.get("coords", None), img_version, tuple(img.get("name", "") for img in task.get("images", []))
    )
    entry = cache.get(t_id)
    if entry is not None and entry["key"] == cache_key:
        return entry["json"]

    export_task = {k: v for k, v in task.items() if not k.startswith("_") and k != "is_new"}
    json_str = json.dumps(export_task, ensure_ascii=False)
    cache[t_id] = {"key": cache_key, "json": json_str}
    return json_str

# --- AUTRES CALLBACKS ---
def delete_task_callback(task_id_to_delete):
    st.session_state.tasks = [t for t in st.session_state.tasks if t["id"] != task_id_to_delete]
    st.session_state.task_json_cache.pop(task_id_to_delete, None)
    recalculer_numeros_reserves()

def add_task_callback(coords=None, batiment="BÂTIMENT A", etage="RDC"):
    new_id = f"new_{time.time_ns()}"
    prochain_num = len(st.session_state.tasks) + 1
    st.session_state.tasks.append({
        "id": new_id,
        "num_reserve": f"RES-{prochain_num:03d}",
        "batiment": batiment,
        "etage": etage,
        "priorite": "Moyenne",
        "categorie": "",
        "titre": "",
        "details": "",
        "commentaire": "",
        "status": "A FAIRE",
        "status_logistique": "En attente date de livraison",
        "entreprise": "",  
        "date_livraison": "Non définie",
        "coords": coords, 
        "images": [{"b64": "", "name": ""}, {"b64": "", "name": ""}, {"b64": "", "name": ""}],
        "is_new": True
    })

def recalculer_numeros_reserves():
    """Garantit une suite de numéros cohérente sans trous suite à une suppression ou un import"""
    for index, task in enumerate(st.session_state.tasks):
        task["num_reserve"] = f"RES-{(index + 1):03d}"

# --- BARRE LATÉRALE : SAUVEGARDE & IMPORT & FILTRES ---
st.sidebar.header("💾 Sauvegarde & Restauration")
uploaded_json = st.sidebar.file_uploader("📂 Importer une sauvegarde JSON (.json)", type=["json"])
uploaded_docx = st.sidebar.file_uploader("📝 Importer depuis un fichier Word (.docx)", type=["docx"])

# CONFIGURATION DES FILTRES DE RAPPORT
st.sidebar.header("🎯 Options de génération")
entreprises_existantes = sorted(list(set([t.get("entreprise", "").strip() for t in st.session_state.tasks if t.get("entreprise", "").strip()])))
filter_entreprise = st.sidebar.selectbox("Filtrer par Entreprise :", ["Toutes les entreprises"] + entreprises_existantes)

categories_existantes = sorted(list(set([t.get("categorie", "").strip() for t in st.session_state.tasks if t.get("categorie", "").strip()])))
filter_categorie = st.sidebar.selectbox("Filtrer par Catégorie :", ["Toutes les catégories"] + categories_existantes)

# LOGIQUE D'IMPORTATION DU FICHIER WORD (.DOCX)
if uploaded_docx is not None and st.sidebar.button("🔄 Charger le fichier Word"):
    try:
        doc = Document(uploaded_docx)
        word_tasks = []
        if doc.tables:
            for table in doc.tables:
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 3:
                        word_tasks.append({
                            "id": str(time.time_ns()) + f"_{len(word_tasks)}",
                            "num_reserve": "", 
                            "batiment": cells[0] if "A" in cells[0] or "B" in cells[0] else "BÂTIMENT B",
                            "etage": "RDC",
                            "priorite": "Moyenne",
                            "categorie": cells[1],
                            "titre": cells[2],
                            "details": cells[3] if len(cells) > 3 else "",
                            "commentaire": "",
                            "status": "A FAIRE",
                            "status_logistique": "En attente date de livraison",
                            "entreprise": "",
                            "date_livraison": "Non définie",
                            "coords": None,
                            "images": [{"b64": "", "name": ""}, {"b64": "", "name": ""}, {"b64": "", "name": ""}]
                        })
        if word_tasks:
            st.session_state.tasks = word_tasks
            recalculer_numeros_reserves()
            st.sidebar.success(f"🎉 {len(word_tasks)} réserves importées !")
    except Exception as e:
        st.sidebar.error(f"Erreur lors de la lecture du Word : {e}")

# LOGIQUE D'IMPORTATION DU FICHIER JSON
if uploaded_json is not None:
    if st.session_state.loaded_json_file_name != uploaded_json.name:
        try:
            loaded_data = json.load(uploaded_json)
            if isinstance(loaded_data, dict) and "tasks" in loaded_data:
                st.session_state.floor_plans = loaded_data.get("plans", {})
                tasks_list = loaded_data["tasks"]
            else:
                tasks_list = loaded_data

            for idx, item in enumerate(tasks_list):
                if "id" not in item:
                    item["id"] = f"imported_{idx}_{time.time_ns()}"
                if "num_reserve" not in item:
                    item["num_reserve"] = ""
                if "etage" not in item:
                    item["etage"] = "RDC"
                if "priorite" not in item:
                    item["priorite"] = "Moyenne"
                if "commentaire" not in item:
                    item["commentaire"] = ""
                if "entreprise" not in item:
                    item["entreprise"] = ""
                if "status" not in item:
                    item["status"] = "A FAIRE"
                if "status_logistique" not in item:
                    item["status_logistique"] = "En attente date de livraison"
                if "date_livraison" not in item:
                    item["date_livraison"] = "Non définie"
                if "coords" not in item:
                    item["coords"] = None
                if "images" not in item:
                    item["images"] = []
                    if item.get("img_b64"):
                        item["images"].append({
                            "b64": item["img_b64"],
                            "name": item.get("file_name", "Photo 1")
                        })
                while len(item["images"]) < 3:
                    item["images"].append({"b64": "", "name": ""})
            st.session_state.tasks = tasks_list
            recalculer_numeros_reserves()
            st.session_state.loaded_json_file_name = uploaded_json.name
            st.sidebar.success("✅ Sauvegarde JSON importée avec succès !")
        except Exception as e:
            st.sidebar.error(f"Erreur lors de l'import JSON : {e}")
else:
    st.session_state.loaded_json_file_name = None

if not st.session_state.tasks:
    st.session_state.tasks = [
        {
            "id": "init_1",
            "num_reserve": "RES-001",
            "batiment": "BÂTIMENT B",
            "etage": "R+1",
            "priorite": "Moyenne",
            "categorie": "Serrurerie & Organigramme",
            "titre": "Pose des cylindres de portes",
            "details": "Intervenant : Mickael\nFinalisation complète selon l'organigramme.",
            "commentaire": "",
            "status": "A FAIRE",
            "status_logistique": "En attente date de livraison",
            "entreprise": "Serrurerie Pro",
            "date_livraison": "Non définie",
            "coords": None,
            "images": [{"b64": "", "name": ""}, {"b64": "", "name": ""}, {"b64": "", "name": ""}]
        }
    ]

# --- MANAGEMENT MULTI-PLANS PAR NIVEAU ---
st.header("🗺️ Plan de repérage par niveau")

col_map_sel1, col_map_sel2 = st.columns(2)
with col_map_sel1:
    plan_active_bat = st.selectbox("Sélectionner le Bâtiment pour le plan", ["BÂTIMENT A", "BÂTIMENT B"], key="plan_active_bat")
with col_map_sel2:
    plan_active_etage = st.text_input("Sélectionner / Écrire l'Étage pour le plan", "RDC", key="plan_active_etage").strip().upper()

plan_key = f"{plan_active_bat} - {plan_active_etage}"

col_up1, col_up2 = st.columns([3, 1])
with col_up1:
    bg_plan_file = st.file_uploader(f"📂 Charger le plan spécifique pour [ {plan_key} ] (PDF, PNG ou JPG)", type=["pdf", "png", "jpg", "jpeg"], key=f"uploader_{plan_key}")
with col_up2:
    pdf_page_num = st.number_input("N° Page du PDF", min_value=1, value=1, step=1)

if bg_plan_file is not None:
    try:
        file_ext = os.path.splitext(bg_plan_file.name)[1].lower()
        plan_img = None

        if file_ext == ".pdf":
            pdf_bytes = bg_plan_file.read()
            doc_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
            max_pages = len(doc_pdf)
            target_page = min(pdf_page_num - 1, max_pages - 1)
            page = doc_pdf[target_page] 
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) 
            img_data = pix.tobytes("png")
            plan_img = Image.open(io.BytesIO(img_data))
            doc_pdf.close()
        else:
            plan_img = Image.open(bg_plan_file)
        
        buffered_plan = io.BytesIO()
        plan_img.save(buffered_plan, format="JPEG")
        plan_b64 = base64.b64encode(buffered_plan.getvalue()).decode("utf-8")
        st.session_state.floor_plans[plan_key] = f"data:image/jpeg;base64,{plan_b64}"
        st.success(f"🎉 Plan enregistré avec succès pour {plan_key} (Page {pdf_page_num}) !")
    except Exception as e:
        st.error(f"Erreur lors du conversion du plan : {e}")

if plan_key in st.session_state.floor_plans:
    plan_data_url = st.session_state.floor_plans[plan_key]
    bounds = [[0, 0], [1000, 1000]]
    
    m = folium.Map(crs="Simple", tiles=None, max_zoom=3, min_zoom=-1, zoom_start=0, location=[500, 500])
    folium.raster_layers.ImageOverlay(image=plan_data_url, bounds=bounds).add_to(m)
    m.fit_bounds(bounds)
    
    for t in st.session_state.tasks:
        if t.get("coords") is not None and t.get("batiment") == plan_active_bat and t.get("etage", "").strip().upper() == plan_active_etage:
            st_raw = t.get('status', 'A FAIRE')
            p_color = "#e14d43" if st_raw == "A FAIRE" else "#22c55e"
            
            raw_num = t["num_reserve"].split("-")[-1] if "-" in t["num_reserve"] else "00"
            
            icon_html = f"""
            <div style="background-color: {p_color}; color: white; border: 1.5px solid #ffffff; border-radius: 50%; width: 26px; height: 26px; display: flex; align-items: center; justify-content: center; font-family: Arial, sans-serif; font-size: 9px; font-weight: bold; box-shadow: 0px 2px 4px rgba(0,0,0,0.3);">
                {raw_num}
            </div>
            """
            
            popup_content = f"""
            <div style="font-family: Arial, sans-serif; font-size: 11px; min-width: 130px; text-align: left;">
                <b>{t['num_reserve']}</b> [ {t.get('entreprise', 'Non affectée')} ]<br/>
                Chantier: {st_raw}<br/>
                Logistique: {t.get('status_logistique', '-')}<br/><br/>
                <a href="#ancre_{t['num_reserve']}" target="_self" style="display: block; width: 100%; text-decoration: none; background-color: #2563eb; color: white; text-align: center; padding: 6px; border-radius: 4px; font-weight: bold;">👁️ Voir la réserve</a>
            </div>
            """
            
            folium.Marker(
                location=t["coords"],
                icon=folium.DivIcon(icon_size=(26, 26), icon_anchor=(13, 13), html=icon_html),
                popup=folium.Popup(popup_content, max_width=200)
            ).add_to(m)
    
    st.write(f"💡 *Plan actif : **{plan_key}**. Cliquez pour implanter une réserve directement à ce niveau.*")
    map_data = st_folium(m, width=1100, height=600, key=f"plan_{plan_key}")
    
    if map_data and map_data.get("last_clicked"):
        click_coords = [map_data["last_clicked"]["lat"], map_data["last_clicked"]["lng"]]
        if not any(t.get("coords") == click_coords for t in st.session_state.tasks):
            add_task_callback(coords=click_coords, batiment=plan_active_bat, etage=plan_active_etage)
            st.rerun()
else:
    st.info(f"🗺️ Aucun plan disponible pour **{plan_key}**. Importez un fichier ci-dessus pour commencer l'implantation.")


# --- STATISTIQUES ET GRAPHIQUES DE PROGRESSION ---
st.header("📊 Tableau de progression de chantier")
total_t = len(st.session_state.tasks)
if total_t > 0:
    nb_af = sum(1 for t in st.session_state.tasks if t.get('status') == 'A FAIRE')
    nb_cmd = sum(1 for t in st.session_state.tasks if t.get('status_logistique') == 'Commandé')
    nb_fab = sum(1 for t in st.session_state.tasks if t.get('status_logistique') == 'En fabrication')
    nb_rec = sum(1 for t in st.session_state.tasks if "Reçu" in str(t.get('status_logistique')))
    nb_f = sum(1 for t in st.session_state.tasks if t.get('status') == 'FAIT')
    
    pct_levee = (nb_f / total_t) * 100
    st.write(f"**Taux de levée global des réserves : {pct_levee:.1f}%**")
    st.progress(pct_levee / 100)
    
    c_kpi1, c_kpi2, c_kpi3, c_kpi4, c_kpi5, c_kpi6 = st.columns(6)
    c_kpi1.metric("Total Suivis", total_t)
    c_kpi2.metric("🔴 À Faire", nb_af)
    c_kpi3.metric("🔵 Commandé", nb_cmd)
    c_kpi4.metric("🟣 En Fabrication", nb_fab)
    c_kpi5.metric("🔷 Reçu Site / Massy", nb_rec)
    c_kpi6.metric("🟢 Levées", nb_f)
else:
    st.info("Aucune statistique disponible pour le moment.")


# --- MOTEUR DE RECHERCHE SECURISÉ MULTI-CHAMPS ---
st.header("🔍 Moteur de recherche")
with st.form(key="search_form", clear_on_submit=False):
    search_query = st.text_input(
        "Tapez un lot, une entreprise, ou un numéro (ex: '003', '3', 'RES-003', ou 'es') :",
        st.session_state.get("last_search_query", "")
    )
    submit_search = st.form_submit_button(label="🔍 Valider la recherche")
    if submit_search:
        st.session_state["last_search_query"] = search_query

search_query_active = st.session_state.get("last_search_query", "").strip()

# --- BLOC DE DÉTECTION INTELLIGENTE DE L'ANCRE ---
target_anchor_id = None
if search_query_active:
    query_clean = search_query_active.upper()
    
    # S'active pour le défilement uniquement si c'est un format de numéro
    if query_clean.startswith("RES-"):
        target_anchor_id = query_clean
    elif query_clean.isdigit():
        try:
            num_int = int(query_clean)
            target_anchor_id = f"RES-{num_int:03d}"
        except:
            pass

# Défilement forcé si un numéro valide a été détecté
if target_anchor_id:
    st.markdown(f"""
        <script>
            setTimeout(function() {{
                var element = window.parent.document.getElementById('ancre_{target_anchor_id}');
                if (element) {{
                    element.scrollIntoView({{behavior: 'smooth', block: 'start'}});
                }}
            }}, 150);
        </script>
    """, unsafe_allow_html=True)


# --- FORMULAIRE DE GESTION ---
st.header("1. Gestion des tâches et photos")

today_date = datetime.date.today()

for i, task in enumerate(st.session_state.tasks):
    if "id" not in task:
        task["id"] = f"fix_{i}_{time.time_ns()}"
    
    if not task.get("num_reserve"):
        task["num_reserve"] = f"RES-{(i + 1):03d}"

    t_id = task["id"]

    if task.get("is_new", False):
        match_found = True
    else:
        # Recherche multi-champs exhaustive (Lots, Artisan, Titre, Numéro, Commentaires)
        search_text = f"{task.get('num_reserve','')} {task.get('batiment','')} {task.get('etage','')} {task.get('priorite','')} {task.get('categorie','')} {task.get('titre','')} {task.get('details','')} {task.get('commentaire','')} {task.get('entreprise','')} {task.get('date_livraison','')}".lower()
        search_terms = search_query_active.lower().split()
        match_found = all(term in search_text for term in search_terms)

    if not match_found:
        continue

    current_live_title = task.get('titre', '').strip()
    if not current_live_title:
        current_live_title = "Nouvelle réserve (sans titre)"

    current_live_num = task.get('num_reserve', 'RES-000')
    current_live_bat = task.get('batiment', 'BÂTIMENT A')
    current_live_etage = task.get('etage', 'RDC')
    current_live_prio = task.get('priorite', 'Moyenne')
    current_live_status = task.get('status', 'A FAIRE')
    current_live_logistique = task.get('status_logistique', 'En attente date de livraison')
    current_ent = task.get('entreprise', '')
    current_cat_display = task.get('categorie', '')
    has_plan_marker = f"📍 Placé sur plan ({current_live_bat} - {current_live_etage})" if task.get("coords") else "⚠️ Hors plan"

    raw_date_str = task.get("date_livraison", "Non définie")
    is_date_undefined = (raw_date_str == "Non définie")
    
    if is_date_undefined:
        initial_calendar_value = today_date
        is_delayed = False
    else:
        try:
            initial_calendar_value = datetime.datetime.strptime(raw_date_str, "%Y-%m-%d").date()
            is_delayed = initial_calendar_value < today_date and current_live_status != "FAIT" and "Reçu" not in current_live_logistique
        except:
            initial_calendar_value = today_date
            is_delayed = False

    delay_tag = " | ⚠️ RETARD LIVRAISON" if is_delayed else ""

    st.markdown(f'<div id="ancre_{current_live_num}"></div>', unsafe_allow_html=True)

    with st.expander(f"🆔 {current_live_num} | 📁 [{current_cat_display if current_cat_display else 'Sans catégorie'}] | 🏢 [{current_ent if current_ent else 'Non assignée'}] {delay_tag}", expanded=True):
        col1, col2 = st.columns([1, 1])

        with col1:
            col_bat, col_etage = st.columns(2)
            with col_bat:
                default_bat = 0 if task.get('batiment') == "BÂTIMENT A" else 1
                t_bat = st.selectbox(f"Bâtiment", ["BÂTIMENT A", "BÂTIMENT B"], index=default_bat, key=f"bat_{t_id}", on_change=st.rerun)
                task['batiment'] = t_bat
            with col_etage:
                t_etage = st.text_input(f"Étage", task.get('etage', 'RDC'), key=f"etage_{t_id}", on_change=st.rerun)
                task['etage'] = t_etage

            t_ent = st.text_input(f"Entreprise / Artisan responsable", task.get('entreprise', ''), key=f"ent_{t_id}", on_change=st.rerun)
            task['entreprise'] = t_ent

            t_cat = st.text_input(f"Catégorie / Lot Logistique (ex: Stores)", task.get('categorie', ''), key=f"cat_{t_id}", on_change=st.rerun)
            if t_cat != task.get('categorie', ''):
                task['is_new'] = False
            task['categorie'] = t_cat

            t_titre = st.text_input(f"Intitulé de la réserve / Élément", task.get('titre', ''), key=f"titre_{t_id}")
            if t_titre != task.get('titre', ''):
                task['is_new'] = False
            task['titre'] = t_titre

        with col2:
            t_details = st.text_area(f"Détails (Effectifs, Intervenants)", task.get('details', ''), key=f"det_{t_id}")
            task['details'] = t_details

            col_stat_ch, col_stat_log = st.columns(2)
            with col_stat_ch:
                status_ch_options = ["A FAIRE", "FAIT"]
                current_ch_status = task.get('status', 'A FAIRE')
                if current_ch_status not in status_ch_options:
                    current_ch_status = "A FAIRE"
                t_status_ch = st.selectbox(f"Statut Chantier", status_ch_options, index=status_ch_options.index(current_ch_status), key=f"stat_{t_id}", on_change=st.rerun)
                task['status'] = t_status_ch

            with col_stat_log:
                status_log_options = ["Commandé", "En fabrication", "Reçu sur site", "Reçu à Massy", "En attente date de livraison"]
                current_log_status = task.get('status_logistique', 'En attente date de livraison')
                if current_log_status not in status_log_options:
                    current_log_status = "En attente date de livraison"
                t_status_log = st.selectbox(f"Statut Logistique", status_log_options, index=status_log_options.index(current_log_status), key=f"stat_log_{t_id}", on_change=st.rerun)
                task['status_logistique'] = t_status_log

            col_prio, col_d1_box = st.columns(2)
            with col_prio:
                prio_options = ["Haute", "Moyenne", "Basse"]
                current_prio = task.get('priorite', 'Moyenne')
                if current_prio not in prio_options:
                    current_prio = "Moyenne"
                t_prio = st.selectbox(f"Priorité", prio_options, index=prio_options.index(current_prio), key=f"prio_{t_id}")
                task['priorite'] = t_prio

            col_d1, col_d2 = st.columns([1.5, 1])
            with col_d2:
                check_undef = st.checkbox("Date non définie", value=is_date_undefined, key=f"undef_{t_id}", on_change=st.rerun)
            with col_d1:
                if check_undef:
                    st.text_input("📅 Date de Livraison Attendue", "Non définie", disabled=True, key=f"dis_date_{t_id}")
                    task['date_livraison'] = "Non définie"
                else:
                    t_date = st.date_input(f"📅 Date de Livraison Attendue", initial_calendar_value, key=f"date_{t_id}", on_change=st.rerun)
                    task['date_livraison'] = str(t_date)

            t_comment = st.text_area(f"Commentaires / Réf Commande", task.get('commentaire', ''), key=f"comment_{t_id}")
            task['commentaire'] = t_comment

        st.write("📷 **Photos associées (Max 3)**")
        img_cols = st.columns(3)

        task_images = task.get('images', [])
        while len(task_images) < 3:
            task_images.append({"b64": "", "name": ""})
        task['images'] = task_images

        for img_idx in range(3):
            with img_cols[img_idx]:
                uploader_key = f"file_{t_id}_{img_idx}"

                st.file_uploader(
                    f"Photo {img_idx + 1}",
                    type=["jpg", "jpeg", "png"],
                    key=uploader_key,
                    on_change=update_photo_callback,
                    args=(t_id, img_idx, uploader_key)
                )

                current_b64 = task['images'][img_idx]["b64"]
                current_name = task['images'][img_idx]["name"]

                if current_b64:
                    try:
                        st.image(base64.b64decode(current_b64), width=150, caption=current_name)
                        st.button(
                            f"🗑️ Supprimer photo {img_idx + 1}",
                            key=f"del_img_{t_id}_{img_idx}",
                            on_click=delete_photo_callback,
                            args=(t_id, img_idx)
                        )
                    except Exception as e:
                        pass

        sub_col1, sub_col2 = st.columns([4, 1])
        with sub_col2:
            st.button("❌ Supprimer la ligne", key=f"del_btn_{t_id}", use_container_width=True, on_click=delete_task_callback, args=(t_id,))

st.button("➕ Ajouter une nouvelle ligne manuellement", on_click=add_task_callback)


# --- EXPORTATION COMPLETE JSON ---
if st.sidebar.button("⚙️ Préparer la sauvegarde JSON", use_container_width=True):
    tasks_to_export = [json.loads(get_task_export_json(t)) for t in st.session_state.tasks]
    compiled_str = json.dumps({"plans": st.session_state.floor_plans, "tasks": tasks_to_export}, ensure_ascii=False)
    st.session_state["ready_b64_json"] = base64.b64encode(compiled_str.encode("utf-8")).decode("utf-8")
    st.sidebar.success("✅ Sauvegarde prête !")

if "ready_b64_json" in st.session_state:
    b64_data = st.session_state["ready_b64_json"]
    html_download_button = f"""
        <a href="data:application/json;base64,{b64_data}" download="sauvegarde_complete_chantier.json" style="text-decoration: none;">
            <button style="width: 100%; background-color: #2e7d32; color: white; border: none; padding: 12px 20px; text-align: center; font-size: 15px; font-weight: bold; border-radius: 4px; cursor: pointer; margin-top: 10px; box-shadow: 0px 4px 6px rgba(0,0,0,0.1);">
                📥 Télécharger sauvegarde_complete_chantier.json
            </button>
        </a>
    """
    st.sidebar.markdown(html_download_button, unsafe_allow_html=True)


# --- OPTIONS DU RAPPORT ---
st.header("2. Édition du document professionnel")

with st.expander("⚙️ Informations d'en-tête du rapport", expanded=False):
    col_a, col_b = st.columns(2)
    with col_a:
        chantier_nom = st.text_input("Nom du chantier / operation", "Opération Bâtiments A & B")
        redacteur = st.text_input("Rédigé par", "Alain Autier")
    with col_b:
        echeance = st.text_input("Objectif de livraison", "Fin de Semaine 29 (S29)")
        reference_doc = st.text_input("Référence document", f"REF-{datetime.date.today().strftime('%Y%m%d')}")

if st.button("🚀 Générer et compiler le rapport final", type="primary"):
    def get_category_icon(cat_name):
        c = cat_name.lower()
        if "store" in c: return "🪟"
        if "serrure" in c or "organigramme" in c: return "🔒"
        if "façade" in c or "facade" in c or "habillage" in c: return "🧱"
        return "📋"

    tasks_to_render = st.session_state.tasks
    if filter_entreprise != "Toutes les entreprises":
        tasks_to_render = [t for t in tasks_to_render if t.get("entreprise", "").strip() == filter_entreprise]
    if filter_categorie != "Toutes les catégories":
        tasks_to_render = [t for t in tasks_to_render if t.get("categorie", "").strip() == filter_categorie]

    total = len(tasks_to_render)
    nb_faire = sum(1 for t in tasks_to_render if t.get('status') == 'A FAIRE')
    nb_f = sum(1 for t in tasks_to_render if t.get('status') == 'FAIT')

    # --- ÉTAPE 1 : DESSINS DE PLANS PROPORTIONNELS ---
    html_plans_section = ""
    
    if st.session_state.floor_plans and total > 0:
        html_plans_section += """
        <table class="building-section-table" style="width: 100%; margin-top: 15px; margin-bottom: 10px;">
            <tr><td class="building-title-text">■ PLANS DE REPÉRAGE GENERAUX</td></tr>
        </table>
        """
        
        for p_key, p_b64_url in st.session_state.floor_plans.items():
            try:
                raw_b64 = p_b64_url.split(",")[-1]
                base_img = Image.open(io.BytesIO(base64.b64decode(raw_b64))).convert("RGB")
                draw = ImageDraw.Draw(base_img)
                
                r = max(12, int(base_img.width * 0.016)) 
                font_size = max(10, int(r * 0.85))
                has_markers_on_this_plan = False
                
                for t in tasks_to_render:
                    if t.get("coords") is not None:
                        t_bat = t.get("batiment", "")
                        t_etage = t.get("etage", "").strip().upper()
                        if f"{t_bat} - {t_etage}" == p_key:
                            has_markers_on_this_plan = True
                            
                            st_raw = t.get('status', 'A FAIRE')
                            color_rgb = (220, 38, 38) if st_raw == "A FAIRE" else (34, 197, 94)
                            
                            folium_y, folium_x = t["coords"]
                            pixel_x = int((folium_x / 1000) * base_img.width)
                            pixel_y = int(((1000 - folium_y) / 1000) * base_img.height)
                            
                            draw.ellipse([pixel_x - r, pixel_y - r, pixel_x + r, pixel_y + r], fill=color_rgb, outline=(255,255,255), width=max(1, int(r*0.12)))
                            
                            short_num = t["num_reserve"].split("-")[-1] if "-" in t["num_reserve"] else "00"
                            draw.text((pixel_x, pixel_y), short_num, fill=(255,255,255), anchor="mm", font_size=font_size)
                
                if has_markers_on_this_plan:
                    buf_final_plan = io.BytesIO()
                    base_img.save(buf_final_plan, format="JPEG", quality=75)
                    compiled_plan_b64 = base64.b64encode(buf_final_plan.getvalue()).decode("utf-8")
                    
                    html_plans_section += f"""
                    <div style="page-break-inside: avoid; margin-bottom: 20px; text-align: center; background-color: #ffffff; padding: 10px; border: 0.5pt solid #cbd5e1; border-radius: 4px;">
                        <div style="font-size: 10.5pt; font-weight: bold; color: #1e293b; margin-bottom: 6px; text-align: left;">📍 Niveau : {p_key}</div>
                        <img src="data:image/jpeg;base64,{compiled_plan_b64}" style="width: 100%; max-height: 480px; object-fit: contain;" />
                    </div>
                    """
            except Exception as plan_err:
                pass

    # --- ÉTAPE 2 : COMPILATION DES LIGNES DES CARTES INDIVIDUELLES ---
    html_rows = ""
    current_bat = ""
    current_cat = ""

    sorted_tasks = sorted(
        tasks_to_render, 
        key=lambda x: (x.get('batiment', ''), x.get('etage', ''), priority_weight(x.get('priorite', 'Moyenne')), x.get('categorie', ''))
    )

    for t in sorted_tasks:
        if t['batiment'] != current_bat:
            current_bat = t['batiment']
            html_rows += f"""
            <table class="building-section-table" style="width: 100%; margin-top: 15px; margin-bottom: 4px; -pdf-keep-with-next: true;">
                <tr><td class="building-title-text">■ {current_bat}</td></tr>
            </table>
            """
            current_cat = ""

        if t['categorie'] != current_cat:
            current_cat = t['categorie']
            html_rows += f"""
            <table class="category-section-table" style="width: 100%; margin-top: 8px; margin-bottom: 6px; -pdf-keep-with-next: true;">
                <tr><td class="category-title-text">■ {current_cat if current_cat else "Non spécifié"}</td></tr>
            </table>
            """

        # --- LOGIQUE PHOTO COMPACTE ---
        images_html_block = ""
        t_images = t.get('images', [])
        valid_images = [img for img in t_images if img.get("b64")]

        if valid_images:
            images_html_block += '<table style="border-spacing: 0; padding: 0; margin: 0; width: 100%;">'
            images_html_block += '<tr>'
            for idx in range(min(2, len(valid_images))):
                images_html_block += f"""
                <td style="width: 50%; padding: 2px; vertical-align: top; text-align: left;">
                    <div class="photo-meta"><span class="photo-filename">{valid_images[idx]['name']}</span></div>
                    <img class="thumbnail-img" src="data:image/jpeg;base64,{valid_images[idx]['b64']}" />
                </td>
                """
            if len(valid_images) == 1:
                images_html_block += '<td style="width: 50%;"></td>'
            images_html_block += '</tr>'
            
            if len(valid_images) > 2:
                images_html_block += f"""
                <tr>
                    <td colspan="2" style="width: 100%; padding: 4px 2px 2px 2px; vertical-align: top; text-align: left;">
                        <div class="photo-meta"><span class="photo-filename">{valid_images[2]['name']}</span></div>
                        <img class="thumbnail-img" src="data:image/jpeg;base64,{valid_images[2]['b64']}" />
                    </td>
                </tr>
                """
            images_html_block += '</table>'
        else:
            images_html_block = '<span class="img-placeholder">Aucun visuel disponible</span>'

        st_ch_raw = t.get('status', 'A FAIRE')
        st_log_raw = t.get('status_logistique', 'En attente date de livraison')
        
        num_circle_color = "#dc2626" if st_ch_raw == "A FAIRE" else "#22c55e"

        comment_html_block = ""
        if t.get('commentaire', '').strip():
            comment_html_block = f"""
            <tr>
                <td colspan="3" style="padding: 10px 12px; background-color: #f0f9ff; border-top: 0.5pt solid #cbd5e1; vertical-align: top;">
                    <div style="font-size: 8.5pt; color: #0369a1; line-height: 1.4;">
                        <b>Suivi :</b> {t['commentaire'].replace(chr(10), '<br/>')}
                    </div>
                </td>
            </tr>
            """

        raw_dl = t.get('date_livraison', 'Non définie')
        if raw_dl == "Non définie":
            formatted_dl = "Non définie"
        else:
            try:
                formatted_dl = datetime.datetime.strptime(raw_dl, "%Y-%m-%d").strftime("%d/%m/%Y")
            except:
                formatted_dl = raw_dl

        display_etage = t.get('etage', 'RDC')
        display_num = t.get('num_reserve', 'RES-000')
        display_title = t['titre'] if t['titre'].strip() else 'Nouvelle réserve'
        display_ent_text = t.get('entreprise', 'Non spécifiée')
        display_details = t['details'].replace(chr(10), '<br/>') if t['details'].strip() else 'Aucun détail fourni'
        
        html_rows += f"""
        <table class="reserve-card-table" style="width: 100%; margin-bottom: 12px; border-spacing: 0; background-color: #ffffff; border: 0.5pt solid #e2e8f0; border-left: 0.5pt solid {num_circle_color}; page-break-inside: avoid;">
            <tr>
                <td class="cell-description" style="width: 48%; padding: 12px; vertical-align: top; text-align: left;">
                    <div style="margin-top: 0px; margin-bottom: 8px;">
                        <table style="border-spacing: 0; margin: 0; padding: 0; width: 100%; margin-bottom: 8px;">
                            <tr>
                                <td style="background-color: {num_circle_color}; color: #ffffff; padding: 4px 8px; font-size: 8.5pt; font-weight: bold; text-align: left;">
                                    {display_num}
                                </td>
                            </tr>
                        </table>
                    </div>
                    <div class="desc-stacked-line" style="padding-top: 4px;"><b>Étage :</b> {display_etage}</div>
                    <div class="desc-stacked-line"><b>Artisan responsable :</b> {display_ent_text}</div>
                    <div class="desc-stacked-line"><b>Livraison attendue :</b> {formatted_dl}</div>
                    <div class="desc-stacked-line" style="font-weight: bold; margin-top:8px; margin-bottom:6px; color:#0f172a; font-size: 10pt;">{display_title}</div>
                    <div style="font-size: 8.5pt; color: #475569; line-height: 1.4; border-top: 0.5pt dashed #e2e8f0; padding-top: 6px; margin-top: 6px;">
                        {display_details}
                    </div>
                </td>
                <td class="cell-photo" style="width: 39%; padding: 12px; border-left: 0.5pt solid #e2e8f0; border-right: 0.5pt solid {num_circle_color}; vertical-align: top; text-align: left;">
                    {images_html_block}
                </td>
                <td class="cell-status" style="width: 13%; padding: 8px 4px; vertical-align: middle; text-align: center;">
                    <table style="border-spacing: 0; margin: 0; padding: 0; width: 100%;">
                        <tr>
                            <td style="background-color: {num_circle_color}; color: white; font-weight: bold; padding: 4px 2px; font-size: 7pt; text-align: center; border-radius: 2px;">
                                {st_ch_raw}
                            </td>
                        </tr>
                        <tr><td style="height: 4px; font-size: 1px; line-height: 1px;">&nbsp;</td></tr>
                        <tr>
                            <td style="background-color: #f1f5f9; color: #334155; border: 0.5pt solid #cbd5e1; padding: 4px 2px; font-size: 6.5pt; font-weight: normal; text-align: center; border-radius: 2px; line-height: 1.1;">
                                {st_log_raw}
                            </td>
                        </tr>
                    </table>
                </td>
            </tr>
            {comment_html_block}
        </table>
        """

    full_html = f"""
    <!DOCTYPE html>
    <html lang="fr">
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: A4; margin-top: 1.6cm; margin-bottom: 1.6cm; margin-left: 1.5cm; margin-right: 1.5cm;
                @frame footer_frame {{ -pdf-frame-content: footerContent; left: 1.5cm; width: 18cm; top: 28.4cm; height: 1.0cm; }}
            }}
            body {{ font-family: Helvetica, Arial, sans-serif; color: #1e293b; font-size: 9.5pt; line-height: 1.45; background-color: #f8fafc; }}
            
            #footerContent {{ font-size: 8pt; color: #94a3b8; border-top: 0.5pt solid #e2e8f0; padding-top: 3pt; }}
            #footerContent table {{ width: 100%; }} #footerContent td {{ font-size: 8pt; color: #94a3b8; }} .footer-right {{ text-align: right; }}
            .main-report-title {{ font-size: 14pt; font-weight: bold; color: #1e293b; text-transform: uppercase; margin-bottom: 8px; letter-spacing: 0.5px; }}
            .meta-bar-table {{ width: 100%; border: 0.5pt solid #cbd5e1; border-spacing: 0; background-color: #ffffff; margin-bottom: 12px; }}
            .meta-bar-cell {{ padding: 6px 10px; font-size: 8.5pt; color: #475569; border-right: 0.5pt solid #cbd5e1; }}
            .meta-bar-cell b {{ color: #0f172a; }} .meta-bar-right {{ border-right: none; }}
            .summary-table {{ width: 100%; border: 0.5pt solid #cbd5e1; border-spacing: 0; background-color: #ffffff; margin-bottom: 12px; }}
            .summary-table td {{ padding: 6px; font-size: 8.5pt; text-align: center; border-right: 0.5pt solid #cbd5e1; border-bottom: 0.5pt solid #cbd5e1; }}
            .summary-header-td {{ background-color: #f8fafc; font-weight: bold; color: #475569; text-transform: uppercase; }}
            .summary-value-td {{ font-size: 11pt; font-weight: bold; }}
            .building-title-text {{ background-color: #1e293b; color: #ffffff; font-size: 10pt; font-weight: bold; padding: 4px 8px; border-radius: 3px; letter-spacing: 0.3px; }}
            .category-title-text {{ background-color: #eff6ff; color: #1e40af; font-size: 9.5pt; font-weight: bold; padding: 5px; border-left: none; }}
            
            .desc-stacked-line {{ font-size: 9pt; margin-top: 0px; margin-bottom: 3px; padding: 0px; color: #334155; line-height: 1.35; }}
            .desc-stacked-line b {{ color: #0f172a; font-weight: bold; }}
            
            .photo-meta {{ font-size: 7.5pt; color: #64748b; margin-bottom: 2px; font-weight: bold; max-width: 105px; overflow: hidden; white-space: nowrap; }}
            .thumbnail-img {{ width: 105px; height: 80px; object-fit: cover; border: 0.5pt solid #cbd5e1; border-radius: 4px; }}
            .img-placeholder {{ display: block; width: 110px; height: 45px; background-color: #f1f5f9; border: 0.5pt solid #e2e8f0; font-size: 8pt; color: #94a3b8; text-align: center; padding-top: 16px; font-style: italic; border-radius: 4px; }}
        </style>
    </head>
    <body>
        <div id="footerContent">
            <table>
                <tr>
                    <td><b>{chantier_nom}</b> &nbsp;|&nbsp; Filtres: {filter_entreprise} / {filter_categorie}</td>
                    <td class="footer-right">Page <pdf:pagenumber /> / <pdf:pagecount /></td>
                </tr>
            </table>
        </div>
        <div class="main-report-title">Suivi Logistique &amp; État des Commentes Chantier</div>
        <table class="meta-bar-table">
            <tr>
                <td class="meta-bar-cell" style="width: 25%;"><b>Document Réf :</b> {reference_doc}</td>
                <td class="meta-bar-cell" style="width: 25%;"><b>Objectif :</b> {echeance}</td>
                <td class="meta-bar-cell" style="width: 25%;"><b>Rédigé par :</b> {redacteur}</td>
                <td class="meta-bar-cell meta-bar-right" style="width: 25%; text-align: right;"><b>Date :</b> {datetime.date.today().strftime('%d/%m/%Y')}</td>
            </tr>
        </table>
        <table class="summary-table">
            <tr>
                <td class="summary-header-td" style="width: 34%;">Éléments suivis</td>
                <td class="summary-header-td" style="width: 33%;">À traiter (Chantier A FAIRE)</td>
                <td class="summary-header-td" style="width: 33%; border-right: none;">Validés / Levés (Chantier FAIT)</td>
            </tr>
            <tr>
                <td class="summary-value-td" style="border-bottom: none;">{total}</td>
                <td class="summary-value-td" style="color: #dc2626; border-bottom: none;">{nb_faire}</td>
                <td class="summary-value-td" style="color: #166534; border-right: none; border-bottom: none;">{nb_f}</td>
            </tr>
        </table>
        
        {html_plans_section}
        
        {html_rows}
    </body>
    </html>
    """

    with open("rapport_dynamique.html", "w", encoding="utf-8") as f:
        f.write(full_html)
    with open("rapport_dynamique.pdf", "wb") as pdf_file:
        pisa.CreatePDF(full_html, dest=pdf_file)

    st.success(f"🎉 Rapport logistique épuré compilé avec succès !")
    filename_clean = f"Rapport_Logistique_{filter_categorie.replace(' ', '_')}"
    with open("rapport_dynamique.pdf", "rb") as pdf_file:
        st.download_button(label="📥 Télécharger le rapport logistique PDF", data=pdf_file, file_name=f"{filename_clean}.pdf", mime="application/pdf")
    with open("rapport_dynamique.html", "rb") as html_file:
        st.download_button(label="🌐 Télécharger la version HTML autonome", data=html_file, file_name=f"{filename_clean}.html", mime="text/html")
